"""Professional Resume Export service (Phase 3C).

Turns an already-saved ``TailoredResume`` into a clean, ATS-friendly,
single-column PDF and DOCX for an actual job application.

Key rules
---------
- This module NEVER calls the AI provider. It formats the already-approved
  structured content stored on the ``TailoredResume`` row.
- The tailored content is never semantically modified; it is only re-laid-out.
- Contact information is reconstructed from the trusted source ``Resume``
  parsed data (``basic_info``), never invented.
- Empty sections are omitted entirely. No photos, charts, skill bars, icons,
  multi-column layouts, or excessive decoration.
"""

import logging
import re
from io import BytesIO
from typing import Any, Dict, List, Optional

import pymupdf as fitz  # PyMuPDF (already a pinned, production dependency)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_HTTP_RE = re.compile(r"^[a-zA-Z][a-zA-Z0-9+.\-]*://")


# =============================================================================
# Filename handling (safe, no path traversal, no unsafe characters)
# =============================================================================
def sanitize_filename_component(value: Optional[str]) -> str:
    """Reduce arbitrary text to a safe filename component.

    Removes path separators, dots, colons and any other unsafe/special
    characters; collapses runs of them into a single underscore.
    """
    text = (value or "").strip()
    text = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return text[:60].rstrip("_")


def build_export_filename(
    job_title: Optional[str] = None,
    company: Optional[str] = None,
    ext: str = "pdf",
) -> str:
    """Build a safe, human-readable export filename.

    Example: ``CareerPilot_Backend_Developer_Acme.pdf``
    """
    extension = (ext or "").lstrip(".").lower()
    extension = extension if extension in ("pdf", "docx") else "pdf"

    parts = [
        p
        for p in (
            sanitize_filename_component(job_title),
            sanitize_filename_component(company),
        )
        if p
    ]
    stem = "_".join(parts) if parts else "CareerPilot_Tailored_Resume"
    if parts:
        stem = f"CareerPilot_{stem}"
    if len(stem) > 100:
        stem = stem[:100].rstrip("_")
    return f"{stem}.{extension}"


# =============================================================================
# Document model (normalised, provider-agnostic)
# =============================================================================
def _clean(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _build_contact(basic_info: Optional[Dict[str, Any]]) -> List[List]:
    """Build contact lines + which of them are clickable URL entries.

    Returns (``lines``, ``links``) where ``links`` is a list of
    ``{"index": <int>, "url": <str>}``. Only trusted parsed contact data is
    used; nothing is ever invented.
    """
    lines: List[str] = []
    links: List[Dict[str, Any]] = []
    bi = basic_info or {}

    email = _clean(bi.get("email"))
    if email:
        lines.append(email)

    phone = _clean(bi.get("phone"))
    if phone:
        lines.append(phone)

    location = _clean(bi.get("location"))
    if location:
        lines.append(location)

    for key in ("linkedin", "github", "portfolio"):
        raw = _clean(bi.get(key))
        if not raw:
            continue
        url = raw
        if not _HTTP_RE.match(url):
            url = "https://" + url
        lines.append(url)
        links.append({"index": len(lines) - 1, "url": url})

    return lines, links


def build_export_document(tailored: Any, resume: Optional[Any]) -> Dict[str, Any]:
    """Build the normalised document model from a TailoredResume + source Resume.

    ``tailored`` is any object exposing ``structured_data`` and
    ``tailored_content`` (e.g. the SQLAlchemy TailoredResume row). ``resume``
    is the source Resume row (may be None); only its parsed ``basic_info`` is
    used for contact information.
    """
    assert tailored is not None, "tailored resume is required"

    structured = (tailored.structured_data or {}) if hasattr(tailored, "structured_data") else {}
    flattened = (tailored.tailored_content or {}) if hasattr(tailored, "tailored_content") else {}

    basic_info = {}
    if resume is not None and hasattr(resume, "parsed_data"):
        parsed = resume.parsed_data or {}
        if isinstance(parsed, dict):
            basic_info = parsed.get("basic_info") or {}

    name = _clean(basic_info.get("name"))
    contact, contact_links = _build_contact(basic_info)

    summary = _clean((structured.get("summary") or {}).get("tailored")) or _clean(
        flattened.get("summary")
    )
    skills = [
        _clean(s) for s in ((structured.get("skills") or {}).get("kept") or flattened.get("skills") or []) if s
    ]

    experience = []
    for item in structured.get("experience") or []:
        title = _clean(item.get("original_title")) if isinstance(item, dict) else ""
        company = _clean(item.get("company")) if isinstance(item, dict) else ""
        bullets = [
            _clean(b)
            for b in (item.get("tailored_bullets") or item.get("original_bullets") or [])
            if b
        ] if isinstance(item, dict) else []
        if title or company or bullets:
            experience.append({"title": title, "company": company, "bullets": bullets})

    projects = []
    for item in structured.get("projects") or []:
        if not isinstance(item, dict):
            continue
        name_p = _clean(item.get("name"))
        description = _clean(item.get("tailored_description")) or _clean(
            item.get("original_description")
        )
        if name_p or description:
            projects.append({"name": name_p, "description": description})

    education = [
        _clean(e) for e in (structured.get("education") or flattened.get("education") or []) if e
    ]
    certifications = [
        _clean(c)
        for c in (structured.get("certifications") or flattened.get("certifications") or [])
        if c
    ]

    return {
        "name": name,
        "contact": contact,
        "contact_links": contact_links,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
        "certifications": certifications,
    }


# =============================================================================
# PDF renderer (PyMuPDF, selectable text, manual layout engine)
# =============================================================================
class _PdfRenderer:
    A4_W = 595.276
    A4_H = 841.890
    ML = 48.0
    MR = 48.0
    MT = 42.0
    MB = 46.0
    BODY = 9.5
    NAME = 20.0
    HEADING = 10.5
    LINE = 1.35
    FONT = "helv"
    BOLD_FONT = "hebo"

    def __init__(self) -> None:
        self.doc = fitz.open()
        self.page = None
        self.y = self.MT

    @property
    def content_width(self) -> float:
        return self.A4_W - self.ML - self.MR

    @property
    def bottom_limit(self) -> float:
        return self.A4_H - self.MB

    def line_h(self, size: float) -> float:
        return size * self.LINE

    def new_page(self) -> None:
        self.page = self.doc.new_page(width=self.A4_W, height=self.A4_H)
        self.y = self.MT

    def remaining(self) -> float:
        return self.bottom_limit - self.y

    def ensure(self, needed: float) -> None:
        if needed > self.remaining() + 0.5:
            self.new_page()

    @staticmethod
    def wrap(text: str, size: float, width: float, font: str = FONT) -> List[str]:
        """Wrap text to lines of at most ``width`` points.

        Long unbreakable tokens (e.g. URLs) are safely char-split so nothing
        overflows the page.
        """
        cleaned = (text or "").replace("\r", " ").replace("\n", " ").strip()
        if not cleaned:
            return []
        words = cleaned.split()
        lines: List[str] = []
        cur = ""

        def char_split(token: str) -> None:
            piece = ""
            for ch in token:
                if fitz.get_text_length(piece + ch, fontname=font, fontsize=size) <= width:
                    piece += ch
                else:
                    if piece:
                        lines.append(piece)
                    piece = ch
            if piece:
                lines.append(piece)

        for word in words:
            candidate = f"{cur} {word}".strip()
            if fitz.get_text_length(candidate, fontname=font, fontsize=size) <= width:
                cur = candidate
                continue
            if cur:
                lines.append(cur)
                cur = ""
            if fitz.get_text_length(word, fontname=font, fontsize=size) <= width:
                cur = word
            else:
                char_split(word)
        if cur:
            lines.append(cur)
        return lines

    def draw_text_line(self, line: str, size: float, bold: bool = False, x: float = None) -> fitz.Rect:
        x = self.ML if x is None else x
        font = self.BOLD_FONT if bold else self.FONT
        self.page.insert_text((x, self.y + size), line, fontname=font, fontsize=size)
        y0, y1 = self.y, self.y + self.line_h(size)
        x1 = x + fitz.get_text_length(line, fontname=font, fontsize=size)
        self.y += self.line_h(size)
        return fitz.Rect(x, y0, x1, y1)

    def draw_paragraph(
        self,
        text: str,
        size: float,
        bold: bool = False,
        x: float = None,
        wrap_width: float = None,
    ) -> List[fitz.Rect]:
        x = self.ML if x is None else x
        wrap_width = self.content_width - (x - self.ML) if wrap_width is None else wrap_width
        font = self.BOLD_FONT if bold else self.FONT
        lines = self.wrap(text, size, wrap_width, font)
        rects = []
        for ln in lines:
            self.ensure(self.line_h(size))
            rects.append(self.draw_text_line(ln, size, bold, x))
        return rects

    def draw_bullets(self, items: List[str], size: float = BODY) -> None:
        indent = 14.0
        x_text = self.ML + indent
        wrap_width = self.content_width - indent
        dot_radius = 1.25
        for item in items:
            lines = self.wrap(item, size, wrap_width)
            if not lines:
                continue
            needed = len(lines) * self.line_h(size)
            # Keep each bullet together on one page when it fits.
            if needed > self.remaining() and self.remaining() < self.line_h(size):
                self.new_page()
            for i, ln in enumerate(lines):
                self.ensure(self.line_h(size))
                if i == 0:
                    cy = self.y + size * 0.55
                    self.page.draw_circle(
                        fitz.Point(self.ML + 4, cy), dot_radius,
                        color=None, fill=(0, 0, 0), width=0,
                    )
                self.draw_text_line(ln, size, bold=False, x=x_text)

    def draw_heading(self, text: str) -> None:
        # Keep-with-next: heading never stranding at the bottom of a page.
        keep = 2 * self.line_h(self.BODY)
        needed = 12.0 + self.line_h(self.HEADING) + 6.0 + keep
        self.ensure(needed)
        self.y += 12.0
        self.draw_text_line(text, self.HEADING, bold=True)
        self.page.draw_line(
            fitz.Point(self.ML, self.y + 2),
            fitz.Point(self.ML + self.content_width, self.y + 2),
            color=(0.55, 0.55, 0.55),
            width=0.6,
        )
        self.y += 6.0


def render_pdf(document: Dict[str, Any]) -> bytes:
    """Render the document model to A4 PDF bytes (selectable text)."""
    r = _PdfRenderer()
    r.new_page()

    name = document.get("name") or ""
    if name:
        r.ensure(r.line_h(r.NAME))
        r.draw_text_line(name, r.NAME, bold=True)
        r.y += 3.0

    contact = document.get("contact") or []
    if contact:
        link_by_index = {c["index"]: c["url"] for c in (document.get("contact_links") or [])}
        for i, item in enumerate(contact):
            rects = r.draw_paragraph(item, r.BODY)
            url = link_by_index.get(i)
            if url and rects:
                region = rects[0]
                for rect in rects[1:]:
                    region |= rect
                try:
                    r.page.insert_link(
                        {"kind": fitz.LINK_URI, "from": region, "uri": url}
                    )
                except Exception:  # noqa: BLE001 - a broken link must never break export
                    logger.warning("Could not add PDF hyperlink for %r", url)
        r.y += 3.0

    if document.get("summary"):
        r.draw_heading("PROFESSIONAL SUMMARY")
        r.draw_paragraph(document["summary"], r.BODY)

    if document.get("skills"):
        r.draw_heading("SKILLS")
        r.draw_paragraph(", ".join(document["skills"]), r.BODY)

    if document.get("experience"):
        r.draw_heading("EXPERIENCE")
        for exp in document["experience"]:
            header = exp["title"]
            if exp["company"]:
                header = f"{header} | {exp['company']}" if header else exp["company"]
            if header:
                r.ensure(r.line_h(r.BODY))
                r.draw_text_line(header, r.BODY + 0.5, bold=True)
            r.draw_bullets(exp["bullets"])
            r.y += 3.0

    if document.get("projects"):
        r.draw_heading("PROJECTS")
        for proj in document["projects"]:
            if proj["name"]:
                r.draw_text_line(proj["name"], r.BODY + 0.5, bold=True)
            if proj["description"]:
                r.draw_paragraph(proj["description"], r.BODY)
            r.y += 3.0

    if document.get("education"):
        r.draw_heading("EDUCATION")
        r.draw_bullets(document["education"])

    if document.get("certifications"):
        r.draw_heading("CERTIFICATIONS")
        r.draw_bullets(document["certifications"])

    return r.doc.tobytes()


# =============================================================================
# DOCX renderer (python-docx, editable, ATS-friendly)
# =============================================================================
def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Add a clickable hyperlink whose visible text is the readable URL."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")  # 10.5pt
    r_pr.append(size)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_bottom_border(paragraph, color: str = "D9D9D9") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def render_docx(document: Dict[str, Any]) -> bytes:
    """Render the document model to DOCX bytes (editable, pdf-equivalent)."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    def heading(text: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        pf = p.paragraph_format
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        _add_bottom_border(p)

    def bullet(text: str) -> None:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(1)

    name = document.get("name") or ""
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(20)
        p.paragraph_format.space_after = Pt(2)

    contact = document.get("contact") or []
    if contact:
        link_by_index = {c["index"]: c["url"] for c in (document.get("contact_links") or [])}
        p = doc.add_paragraph()
        for i, item in enumerate(contact):
            if i:
                sep = p.add_run("  |  ")
                sep.font.size = Pt(10)
            url = link_by_index.get(i)
            if url:
                _add_hyperlink(p, item, url)
            else:
                run = p.add_run(item)
                run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(10)

    if document.get("summary"):
        heading("PROFESSIONAL SUMMARY")
        p = doc.add_paragraph()
        p.add_run(document["summary"])

    if document.get("skills"):
        heading("SKILLS")
        p = doc.add_paragraph()
        p.add_run(", ".join(document["skills"]))

    if document.get("experience"):
        heading("EXPERIENCE")
        for exp in document["experience"]:
            title = exp["title"]
            if exp["company"]:
                title = f"{title} | {exp['company']}" if title else exp["company"]
            if title:
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
            for b in exp["bullets"]:
                bullet(b)

    if document.get("projects"):
        heading("PROJECTS")
        for proj in document["projects"]:
            if proj["name"]:
                p = doc.add_paragraph()
                run = p.add_run(proj["name"])
                run.bold = True
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
            if proj["description"]:
                doc.add_paragraph(proj["description"])

    if document.get("education"):
        heading("EDUCATION")
        for entry in document["education"]:
            bullet(entry)

    if document.get("certifications"):
        heading("CERTIFICATIONS")
        for cert in document["certifications"]:
            bullet(cert)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# Convenience entry points used by the API
# =============================================================================
def generate_pdf(tailored: Any, resume: Optional[Any]) -> bytes:
    document = build_export_document(tailored, resume)
    return render_pdf(document)


def generate_docx(tailored: Any, resume: Optional[Any]) -> bytes:
    document = build_export_document(tailored, resume)
    return render_docx(document)